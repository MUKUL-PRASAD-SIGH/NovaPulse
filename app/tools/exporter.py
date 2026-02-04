"""Multi-format data exporter V3 - JSON, Markdown, CSV, Word, PDF with Intelligence format."""
import json
import os
import csv
from datetime import datetime
from typing import Dict, Any

# Optional imports for Word and PDF
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


OUTPUT_DIR = "output"


def export_data(
    data: Dict[str, Any],
    filename: str = "report",
    format: str = "json"
) -> str:
    """
    Export data to specified format.
    
    Args:
        data: Dict with news, summary, sentiment, trends
        filename: Base filename (without extension)
        format: 'json', 'markdown', 'csv', 'docx', or 'pdf'
    
    Returns:
        Path to saved file
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if format == "markdown":
        return _export_markdown(data, filename, timestamp)
    elif format == "csv":
        return _export_csv(data, filename, timestamp)
    elif format == "docx":
        return _export_docx(data, filename, timestamp)
    elif format == "pdf":
        return _export_pdf(data, filename, timestamp)
    else:
        return _export_json(data, filename, timestamp)


def _export_json(data: Dict, filename: str, timestamp: str) -> str:
    path = f"{OUTPUT_DIR}/{filename}_{timestamp}.json"
    output = {
        "generated_at": datetime.now().isoformat(),
        "report_type": "Nova Intelligence Report",
        **data
    }
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(output, f, indent=2, ensure_ascii=False)
    return path


def _export_markdown(data: Dict, filename: str, timestamp: str) -> str:
    """Export as professional intelligence-style Markdown."""
    path = f"{OUTPUT_DIR}/{filename}_{timestamp}.md"
    lines = [
        "# 🧠 Nova Intelligence Report",
        f"\n*Generated: {timestamp}*\n",
        "---",
    ]
    
    # Summary Section
    if "summary" in data and data["summary"]:
        s = data["summary"]
        lines.append("\n## 📝 Executive Summary\n")
        summary_text = s.get("summary", str(s)) if isinstance(s, dict) else str(s)
        lines.append(summary_text)
        
        if isinstance(s, dict) and s.get("key_points"):
            lines.append("\n**Key Takeaways:**")
            for point in s.get("key_points", []):
                lines.append(f"- {point}")
        lines.append("")
    
    # Sentiment Intelligence V2
    if "sentiment" in data and data["sentiment"]:
        s = data["sentiment"]
        lines.append("\n## 💭 Sentiment Intelligence\n")
        
        # Mood & Direction
        mood = s.get("mood_label", s.get("overall", "N/A"))
        direction = s.get("direction", "stable")
        direction_icon = "📈" if direction == "improving" else "📉" if direction == "deteriorating" else "➡️"
        lines.append(f"**Mood:** {mood}")
        lines.append(f"**Direction:** {direction_icon} {direction.capitalize()}")
        lines.append("")
        
        # Market Indicators
        momentum = s.get("momentum_strength", "moderate")
        bias = s.get("market_bias", "balanced")
        bias_display = "🟢 Risk-On" if bias == "risk_on" else "🔴 Risk-Off" if bias == "risk_off" else "⚪ Balanced"
        lines.append(f"| Indicator | Value |")
        lines.append(f"|-----------|-------|")
        lines.append(f"| Momentum | {momentum.capitalize()} |")
        lines.append(f"| Market Bias | {bias_display} |")
        lines.append(f"| Confidence | {s.get('confidence', 'medium').capitalize()} |")
        lines.append(f"| Risk Level | {s.get('risk_level', 'low').capitalize()} |")
        lines.append("")
        
        # Analyst Reasoning
        if s.get("reasoning"):
            lines.append(f"**Analyst View:** {s.get('reasoning')}")
            lines.append("")
        
        # Signals
        if s.get("positive_signals"):
            lines.append("**✅ Bullish Signals:**")
            for sig in s.get("positive_signals", []):
                lines.append(f"- {sig}")
            lines.append("")
        
        if s.get("negative_signals"):
            lines.append("**⚠️ Risk Signals:**")
            for sig in s.get("negative_signals", []):
                lines.append(f"- {sig}")
            lines.append("")
        
        # Emerging Themes
        if s.get("emerging_themes"):
            themes = ", ".join(s.get("emerging_themes", []))
            lines.append(f"**🔥 Emerging Themes:** {themes}")
            lines.append("")
    
    # Trends Section V2
    if "trends" in data and data["trends"]:
        t = data["trends"]
        if t.get("trending_topics"):
            lines.append("\n## 📊 Trending Topics\n")
            
            # Rising topics
            if t.get("rising_topics"):
                lines.append("**🔥 Rising:**")
                for topic in t.get("rising_topics", [])[:4]:
                    icon = topic.get("velocity_icon", "📈")
                    lines.append(f"- {icon} {topic.get('topic')} (score: {topic.get('score', 0)})")
                lines.append("")
            
            # Main trends table
            lines.append("| Topic | Score | Velocity |")
            lines.append("|-------|-------|----------|")
            for topic in t.get("trending_topics", [])[:8]:
                icon = topic.get("velocity_icon", "➡️")
                lines.append(f"| {topic.get('topic')} | {topic.get('score', topic.get('mentions', 0))} | {icon} |")
            lines.append("")
            
            # Fading topics
            if t.get("fading_topics"):
                lines.append("**📉 Fading:**")
                for topic in t.get("fading_topics", [])[:3]:
                    lines.append(f"- {topic.get('topic')}")
                lines.append("")
    
    # Articles Section
    if "news" in data and data["news"]:
        lines.append("\n## 📰 Source Articles\n")
        for i, item in enumerate(data["news"][:10], 1):
            source = item.get('source', 'unknown').upper()
            lines.append(f"{i}. [{item.get('title', 'No title')}]({item.get('link', '#')}) `{source}`")
        lines.append("")
    
    # Footer
    lines.append("\n---")
    lines.append("*Report generated by Nova Intelligence Agent*")
    
    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    return path


def _export_csv(data: Dict, filename: str, timestamp: str) -> str:
    path = f"{OUTPUT_DIR}/{filename}_{timestamp}.csv"
    news = data.get("news", [{"title": "No data", "link": "", "source": ""}])
    with open(path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=["title", "link", "source", "published"], extrasaction='ignore')
        writer.writeheader()
        writer.writerows(news)
    return path


def _export_docx(data: Dict, filename: str, timestamp: str) -> str:
    """Export as formatted Word document."""
    path = f"{OUTPUT_DIR}/{filename}_{timestamp}.docx"
    
    if not HAS_DOCX:
        # Fallback to markdown if docx not installed
        return _export_markdown(data, filename, timestamp)
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Nova Intelligence Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph(f"Generated: {timestamp}")
    doc.add_paragraph("─" * 50)
    
    # Executive Summary
    if "summary" in data and data["summary"]:
        s = data["summary"]
        doc.add_heading('Executive Summary', level=1)
        summary_text = s.get("summary", str(s)) if isinstance(s, dict) else str(s)
        doc.add_paragraph(summary_text)
        
        if isinstance(s, dict) and s.get("key_points"):
            doc.add_paragraph("Key Takeaways:", style='Intense Quote')
            for point in s.get("key_points", []):
                doc.add_paragraph(f"• {point}", style='List Bullet')
    
    # Sentiment Intelligence
    if "sentiment" in data and data["sentiment"]:
        s = data["sentiment"]
        doc.add_heading('Sentiment Intelligence', level=1)
        
        mood = s.get("mood_label", s.get("overall", "N/A"))
        direction = s.get("direction", "stable")
        
        p = doc.add_paragraph()
        p.add_run("Mood: ").bold = True
        p.add_run(str(mood))
        
        p = doc.add_paragraph()
        p.add_run("Direction: ").bold = True
        p.add_run(direction.capitalize())
        
        # Market table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        cells = [
            ("Indicator", "Value"),
            ("Momentum", s.get("momentum_strength", "moderate").capitalize()),
            ("Market Bias", s.get("market_bias", "balanced").capitalize()),
            ("Confidence", s.get("confidence", "medium").capitalize()),
            ("Risk Level", s.get("risk_level", "low").capitalize()),
        ]
        for i, (col1, col2) in enumerate(cells):
            table.rows[i].cells[0].text = col1
            table.rows[i].cells[1].text = col2
        
        if s.get("reasoning"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Analyst View: ").bold = True
            p.add_run(s.get("reasoning"))
        
        if s.get("positive_signals"):
            doc.add_paragraph("Bullish Signals:", style='Intense Quote')
            for sig in s.get("positive_signals", []):
                doc.add_paragraph(f"✓ {sig}", style='List Bullet')
        
        if s.get("negative_signals"):
            doc.add_paragraph("Risk Signals:", style='Intense Quote')
            for sig in s.get("negative_signals", []):
                doc.add_paragraph(f"⚠ {sig}", style='List Bullet')
    
    # Trends
    if "trends" in data and data["trends"]:
        t = data["trends"]
        if t.get("trending_topics"):
            doc.add_heading('Trending Topics', level=1)
            table = doc.add_table(rows=len(t["trending_topics"][:8]) + 1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = "Topic"
            table.rows[0].cells[1].text = "Mentions"
            for i, topic in enumerate(t["trending_topics"][:8], 1):
                table.rows[i].cells[0].text = topic.get("topic", "")
                table.rows[i].cells[1].text = str(topic.get("mentions", 0))
    
    # News Articles
    if "news" in data and data["news"]:
        doc.add_heading('Source Articles', level=1)
        for i, item in enumerate(data["news"][:10], 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {item.get('title', 'No title')}").bold = True
            doc.add_paragraph(f"   Source: {item.get('source', 'unknown').upper()}")
            doc.add_paragraph(f"   Link: {item.get('link', '#')}")
    
    # Footer
    doc.add_paragraph("─" * 50)
    doc.add_paragraph("Report generated by Nova Intelligence Agent")
    
    doc.save(path)
    return path


def _export_pdf(data: Dict, filename: str, timestamp: str) -> str:
    """Export as formatted PDF document."""
    path = f"{OUTPUT_DIR}/{filename}_{timestamp}.pdf"
    
    if not HAS_PDF:
        # Fallback to markdown if reportlab not installed
        return _export_markdown(data, filename, timestamp)
    
    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, spaceAfter=20, alignment=1)
    heading_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=16, spaceAfter=12, textColor=colors.HexColor('#6366f1'))
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=11, spaceAfter=8)
    
    # Title
    story.append(Paragraph("Nova Intelligence Report", title_style))
    story.append(Paragraph(f"Generated: {timestamp}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Executive Summary
    if "summary" in data and data["summary"]:
        s = data["summary"]
        story.append(Paragraph("Executive Summary", heading_style))
        summary_text = s.get("summary", str(s)) if isinstance(s, dict) else str(s)
        story.append(Paragraph(summary_text, body_style))
        story.append(Spacer(1, 12))
    
    # Sentiment Intelligence
    if "sentiment" in data and data["sentiment"]:
        s = data["sentiment"]
        story.append(Paragraph("Sentiment Intelligence", heading_style))
        
        mood = s.get("mood_label", s.get("overall", "N/A"))
        direction = s.get("direction", "stable")
        
        story.append(Paragraph(f"<b>Mood:</b> {mood}", body_style))
        story.append(Paragraph(f"<b>Direction:</b> {direction.capitalize()}", body_style))
        
        # Market table
        table_data = [
            ["Indicator", "Value"],
            ["Momentum", s.get("momentum_strength", "moderate").capitalize()],
            ["Market Bias", s.get("market_bias", "balanced").capitalize()],
            ["Confidence", s.get("confidence", "medium").capitalize()],
            ["Risk Level", s.get("risk_level", "low").capitalize()],
        ]
        t = Table(table_data, colWidths=[2*inch, 2*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6366f1')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f3f4f6')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))
        
        if s.get("reasoning"):
            story.append(Paragraph(f"<b>Analyst View:</b> {s.get('reasoning')}", body_style))
    
    # Trends
    if "trends" in data and data["trends"]:
        t = data["trends"]
        if t.get("trending_topics"):
            story.append(Paragraph("Trending Topics", heading_style))
            table_data = [["Topic", "Mentions"]]
            for topic in t["trending_topics"][:8]:
                table_data.append([topic.get("topic", ""), str(topic.get("mentions", 0))])
            
            trend_table = Table(table_data, colWidths=[3*inch, 1*inch])
            trend_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#10b981')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
            ]))
            story.append(trend_table)
            story.append(Spacer(1, 12))
    
    # News Articles
    if "news" in data and data["news"]:
        story.append(Paragraph("Source Articles", heading_style))
        for i, item in enumerate(data["news"][:10], 1):
            story.append(Paragraph(f"<b>{i}. {item.get('title', 'No title')}</b>", body_style))
            story.append(Paragraph(f"Source: {item.get('source', 'unknown').upper()}", styles['Normal']))
    
    # Footer
    story.append(Spacer(1, 30))
    story.append(Paragraph("Report generated by Nova Intelligence Agent", styles['Normal']))
    
    doc.build(story)
    return path

